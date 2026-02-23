from pathlib import Path

from docx import Document

from src.services.document_processors.docx_processor import DOCXProcessor


def test_docx_extraction(tmp_path: Path):
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Heading", level=1)
    doc.add_paragraph("Body text")
    doc.save(path)

    processor = DOCXProcessor()
    ok, _ = processor.validate_docx(path)
    assert ok
    pages = processor.extract_text_with_metadata(path)
    assert pages
    assert "Heading" in pages[0]["content"]


def test_validate_doc_rejected(tmp_path: Path):
    path = tmp_path / "legacy.doc"
    path.write_text("dummy", encoding="utf-8")
    processor = DOCXProcessor()
    ok, message = processor.validate_docx(path)
    assert not ok
    assert ".docx" in message
