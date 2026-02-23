"""DOCX extraction utilities."""

from __future__ import annotations

from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document

from src.utils.logger import get_logger


class DOCXProcessor:
    """Extract text, heading structure, and metadata from Word documents."""

    def __init__(self) -> None:
        self.supported_extensions = [".docx", ".doc"]
        self.logger = get_logger(__name__)

    def extract_text_with_metadata(self, file_path: Path) -> List[Dict]:
        """Extract paragraph and table text as a single logical page."""
        doc = Document(str(file_path))
        lines: List[str] = []

        for paragraph in doc.paragraphs:
            text, is_heading = self._process_paragraph(paragraph)
            if not text:
                continue
            if is_heading:
                lines.append(f"## {text}")
            else:
                lines.append(text)

        table_text = self._tables_to_text(doc.tables)
        if table_text:
            lines.append("\n[TABLES]")
            lines.append(table_text)

        return [
            {
                "page_number": 1,
                "content": "\n".join(lines).strip(),
                "tables": table_text,
                "has_images": False,
                "headers_footers": self._extract_headers_footers(doc),
            }
        ]

    def _process_paragraph(self, paragraph) -> Tuple[str, bool]:
        """Return normalized paragraph text and heading flag."""
        text = (paragraph.text or "").strip()
        if not text:
            return "", False
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        is_heading = "heading" in style_name
        return text, is_heading

    def _tables_to_text(self, tables) -> str:
        """Convert Word tables to pipe-delimited text."""
        lines: List[str] = []
        for table in tables or []:
            for row in table.rows:
                cells = [(cell.text or "").strip().replace("\n", " ") for cell in row.cells]
                lines.append(" | ".join(cells))
            lines.append("")
        return "\n".join(lines).strip()

    def _extract_headers_footers(self, doc) -> Dict:
        """Extract header and footer text from all sections."""
        headers: List[str] = []
        footers: List[str] = []
        for section in doc.sections:
            header_text = " ".join(
                p.text.strip() for p in section.header.paragraphs if p.text.strip()
            )
            footer_text = " ".join(
                p.text.strip() for p in section.footer.paragraphs if p.text.strip()
            )
            if header_text:
                headers.append(header_text)
            if footer_text:
                footers.append(footer_text)
        return {"headers": headers, "footers": footers}

    def get_document_metadata(self, file_path: Path) -> Dict:
        """Extract DOCX core properties."""
        doc = Document(str(file_path))
        props = doc.core_properties
        word_count = sum(len((p.text or "").split()) for p in doc.paragraphs)
        return {
            "title": props.title,
            "author": props.author,
            "subject": props.subject,
            "created": props.created.isoformat() if props.created else None,
            "modified": props.modified.isoformat() if props.modified else None,
            "word_count": word_count,
        }

    def validate_docx(self, file_path: Path) -> Tuple[bool, str]:
        """Validate DOCX readability."""
        if not file_path.exists():
            return False, "File not found."
        ext = file_path.suffix.lower()
        if ext == ".doc":
            return False, "Legacy .doc format is not supported. Convert to .docx."
        if ext != ".docx":
            return False, "Unsupported extension. Expected .docx."
        try:
            _ = Document(str(file_path))
        except Exception as exc:
            return False, f"Invalid DOCX: {exc}"
        return True, ""
